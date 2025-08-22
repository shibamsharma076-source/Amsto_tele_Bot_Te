import os
import io
from flask import Flask, render_template, request, redirect, url_for, flash, session , Response, send_file
import sqlite3
from datetime import datetime, timedelta
from werkzeug.security import generate_password_hash
import csv, json
import pandas as pd
from docx import Document


app = Flask(__name__)
app.secret_key = "supersecretkey"  # change to secure key
# app.secret_key = os.environ.get("FLASK_SECRET_KEY")
# if not app.secret_key:
#     raise RuntimeError("FLASK_SECRET_KEY environment variable not set.")


@app.route("/create_admin", methods=["GET", "POST"])
def create_admin():
    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")

        if not username or not password:
            flash("⚠ Username and password required", "danger")
            return redirect(url_for("create_admin"))

        hashed_pw = generate_password_hash(password)

        conn = sqlite3.connect("db/bot_users.db")
        cur = conn.cursor()
        try:
            cur.execute("INSERT INTO admins (username, password) VALUES (?, ?)", (username, hashed_pw))
            conn.commit()
            flash("✅ Admin created successfully!", "success")
        except sqlite3.IntegrityError:
            flash("⚠ Username already exists", "warning")
            return redirect(url_for("create_admin"))
        finally:
            conn.close()

        return redirect(url_for("dashboard"))

    return render_template("create_admin.html")







# ---- DB Helper ----
def query_db(db_path, query, args=(), one=False):
    conn = sqlite3.connect(db_path)
    cur = conn.cursor()
    cur.execute(query, args)
    rv = cur.fetchall()
    conn.close()
    return (rv[0] if rv else None) if one else rv

# ---- Admin Login ----
@app.route("/", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]

        # simple check (replace with DB check later)
        if username == "admin" and password == "admin123":
            session["admin"] = True
            return redirect(url_for("dashboard"))
        else:
            return render_template("login.html", error="Invalid credentials")
    return render_template("login.html")

@app.route("/logout")
def logout():
    session.pop("admin", None)
    return redirect(url_for("login"))





# Delete a Streamlit user
@app.route("/delete_streamlit_user/<int:user_id>")
def delete_streamlit_user(user_id):
    if "admin" not in session:
        return redirect(url_for("login"))

    conn = sqlite3.connect("db/users.db")
    cur = conn.cursor()
    cur.execute("DELETE FROM users WHERE id=?", (user_id,))
    conn.commit()
    conn.close()

    return redirect(url_for("dashboard"))






# ---- Delete Telegram User ----
@app.route("/delete_telegram_user/<int:telegram_id>", methods=["POST"])
def delete_telegram_user(telegram_id):
    if "admin" not in session:
        return redirect(url_for("login"))
    conn = sqlite3.connect("db/bot_users.db")
    cur = conn.cursor()
    cur.execute("DELETE FROM users WHERE telegram_id = ?", (telegram_id,))
    conn.commit()
    conn.close()
    return redirect(url_for("dashboard"))

# ---- Delete Error ----
@app.route("/delete_error/<int:error_id>", methods=["POST"])
def delete_error(error_id):
    if "admin" not in session:
        return redirect(url_for("login"))
    conn = sqlite3.connect("db/bot_errors.db")
    cur = conn.cursor()
    cur.execute("DELETE FROM errors WHERE id = ?", (error_id,))
    conn.commit()
    conn.close()
    return redirect(url_for("dashboard"))









@app.route("/all_data/<string:table>")
def all_data(table):
    if table not in ["errors", "users"]:
        return "Invalid table", 400
    rows = query_db("db/bot_users.db", f"SELECT * FROM {table} ORDER BY id DESC")
    return render_template("all_data.html", rows=rows, table=table)





@app.route("/export/<table>/<format>")
def export_data(table, format):
    valid_tables = ["errors", "users"]
    valid_formats = ["csv", "json", "excel", "docx", "sql"]

    if table not in valid_tables or format not in valid_formats:
        return "Invalid request", 400

    conn = sqlite3.connect("db/bot_users.db")
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute(f"SELECT * FROM {table}")
    rows = cur.fetchall()
    conn.close()

    if not rows:
        return f"No data in {table}", 404

    df = pd.DataFrame([dict(r) for r in rows])

    # --- Export formats ---
    if format == "csv":
        return Response(
            df.to_csv(index=False),
            mimetype="text/csv",
            headers={"Content-disposition": f"attachment; filename={table}.csv"}
        )

    elif format == "json":
        return Response(
            df.to_json(orient="records"),
            mimetype="application/json",
            headers={"Content-disposition": f"attachment; filename={table}.json"}
        )

    elif format == "excel":
        mem = io.BytesIO()
        with pd.ExcelWriter(mem, engine="xlsxwriter") as writer:
            df.to_excel(writer, index=False)
        mem.seek(0)
        return send_file(mem,
                         as_attachment=True,
                         download_name=f"{table}.xlsx",
                         mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    elif format == "docx":
        from docx import Document
        doc = Document()
        doc.add_heading(f"{table} Export", 0)
        for row in rows:
            doc.add_paragraph(str(dict(row)))
        mem = io.BytesIO()
        doc.save(mem)
        mem.seek(0)
        return send_file(mem,
                         as_attachment=True,
                         download_name=f"{table}.docx",
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    elif format == "sql":
        sql_dump = "\n".join(
            [f"INSERT INTO {table} ({', '.join(df.columns)}) VALUES ({', '.join([repr(v) for v in row])});"
             for row in df.values.tolist()]
        )
        return Response(
            sql_dump,
            mimetype="text/plain",
            headers={"Content-disposition": f"attachment; filename={table}.sql"}
        )






















# ---- Dashboard ----
@app.route("/dashboard")
def dashboard():
    if "admin" not in session:
        return redirect(url_for("login"))

    # ---- Web App Users (users.db) ----
    total_users = query_db("db/users.db", "SELECT COUNT(*) FROM users", one=True)[0]

    last_month = (datetime.now() - timedelta(days=30)).isoformat()
    new_users = query_db("db/users.db", "SELECT COUNT(*) FROM users WHERE created_at >= ?", (last_month,), one=True)[0]

    monthly_users = query_db("db/users.db",
        "SELECT strftime('%Y-%m', created_at), COUNT(*) FROM users GROUP BY strftime('%Y-%m', created_at) ORDER BY 1"
    )

    # ---- Streamlit Registered Users ----
    # streamlit_users = query_db("db/users.db",
    #     "SELECT id, full_name, dob, username, email, role, created_at FROM users ORDER BY created_at DESC LIMIT 20"
    # )
    streamlit_users = query_db("db/users.db",
        "SELECT id, full_name, dob, username, email, password_hash, role, created_at FROM users ORDER BY created_at DESC LIMIT 20"
    )




    # ---- Telegram Bot Users (bot_users.db) ----
    telegram_users = query_db("db/bot_users.db",
        "SELECT telegram_id, username, first_name, last_name, phone, email, latitude, longitude, created_at FROM users ORDER BY created_at DESC"
    )


    # Only latest 20 errors and users
    recent_errors = query_db("db/bot_users.db", "SELECT * FROM errors ORDER BY id DESC LIMIT 20")
    recent_telegram_users = query_db("db/bot_users.db", "SELECT * FROM users ORDER BY id DESC LIMIT 20")

    # ---- Errors (bot_errors.db) ----
    try:
        errors = query_db("db/bot_errors.db", "SELECT id, error_message, created_at FROM errors ORDER BY created_at DESC LIMIT 20")
    except Exception:
        errors = []

    return render_template("dashboard.html",
                           total_users=total_users,
                           new_users=new_users,
                           monthly_users=monthly_users,streamlit_users=streamlit_users,
                           telegram_users=telegram_users,
                           errors=errors,
                           recent_errors=recent_errors,
                           recent_telegram_users=recent_telegram_users)

if __name__ == "__main__":
    app.run(port=5001, debug=True)
